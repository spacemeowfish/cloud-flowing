from __future__ import annotations

from pathlib import Path
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(r"D:\my new work\cloud flowing")
ASSET_DIR = ROOT / "work" / "mindmap_assets"
OUTPUT = ROOT / "端云协同AI终端四层功能思维导图.docx"

FONT_REGULAR = r"C:\Windows\Fonts\msyh.ttc"
FONT_BOLD = r"C:\Windows\Fonts\msyhbd.ttc"

CANVAS = (2600, 1460)
BG = "#F7F9FC"
INK = "#17212B"
MUTED = "#66727F"
CENTRAL = "#263746"
FEASIBILITY = {
    "高": ("#2E7D5B", "#EAF6F0"),
    "中": ("#B7791F", "#FFF6E3"),
    "低": ("#C14953", "#FDECEF"),
}
GROUP_COLORS = [
    ("#2F6F9F", "#EAF3FA"),
    ("#287F7B", "#E8F6F4"),
    ("#675AA8", "#F0EEFA"),
    ("#A85E3B", "#FBEEE8"),
    ("#4F7D45", "#EDF6EA"),
    ("#9B762B", "#FBF4E4"),
    ("#8B5678", "#F8EDF4"),
    ("#526675", "#EEF2F5"),
]


MAPS = {
    "架构层": [
        ("交互入口", [("语音输入", "L3", "中"), ("触屏交互", "L2", "高"), ("PC/手机接入", "L3", "中")]),
        ("本地编排", [("意图识别", "L2", "高"), ("参数提取", "L3", "中"), ("工具调度", "L3", "中")]),
        ("工具与连接器", [("本地文件工具", "L2", "高"), ("业务系统连接", "L4", "中"), ("PC端Agent", "L3", "中")]),
        ("端云路由", [("本地/云端判定", "L3", "中"), ("云端能力兜底", "L2", "高"), ("断网降级", "L3", "中")]),
        ("数据与权限", [("身份与权限", "L4", "中"), ("数据隔离", "L4", "中"), ("操作审计", "L3", "高")]),
        ("任务与容错", [("任务状态机", "L3", "高"), ("重试与幂等", "L4", "中"), ("离线任务队列", "L3", "中")]),
    ],
    "应用层": [
        ("语音交互", [("唤醒词", "L2", "高"), ("语音转文字", "L2", "高"), ("文字转语音", "L2", "高")]),
        ("屏幕交互", [("状态展示", "L2", "高"), ("确认/取消", "L2", "高"), ("异常提示", "L2", "高")]),
        ("多端联动", [("PC任务同步", "L3", "中"), ("手机状态查看", "L3", "中"), ("通知同步", "L3", "中")]),
        ("任务管理", [("任务创建", "L2", "高"), ("进度跟踪", "L2", "高"), ("历史记录", "L2", "高")]),
        ("个性化设置", [("偏好设置", "L2", "高"), ("语气风格", "L2", "高"), ("静默模式", "L1", "高")]),
        ("系统状态反馈", [("网络状态", "L1", "高"), ("设备健康", "L2", "高"), ("端/云状态标识", "L2", "高")]),
    ],
    "生活层": [
        ("提醒与日程", [("闹钟提醒", "L1", "高"), ("日程管理", "L2", "高"), ("待办追踪", "L2", "高")]),
        ("信息查询", [("天气查询", "L1", "高"), ("资讯摘要", "L2", "高"), ("生活百科", "L2", "高")]),
        ("生活记录", [("语音备忘", "L2", "高"), ("习惯记录", "L2", "高"), ("健康记录", "L3", "中")]),
        ("出行辅助", [("行程查询", "L2", "高"), ("路线建议", "L2", "高"), ("订票引导", "L4", "中")]),
        ("家庭共享", [("家庭提醒", "L3", "中"), ("共享日历", "L3", "中"), ("多用户区分", "L4", "中")]),
        ("娱乐陪伴", [("音乐播放", "L2", "高"), ("故事笑话", "L2", "高"), ("开放式闲聊", "L2", "高")]),
        ("智能家居协同", [("设备状态查询", "L3", "中"), ("家居场景控制", "L4", "中"), ("跨生态协同", "L5", "低")]),
    ],
    "工作层": [
        ("管理简报", [("每日简报", "L3", "中"), ("重点事项提炼", "L3", "中"), ("风险摘要", "L4", "中")]),
        ("项目进度", [("状态查询", "L3", "中"), ("里程碑提醒", "L2", "高"), ("异常追踪", "L3", "中")]),
        ("会议辅助", [("会议记录", "L3", "高"), ("纪要生成", "L3", "高"), ("行动项提取", "L3", "中")]),
        ("文件与报表", [("文件查找", "L2", "高"), ("报表打开", "L2", "高"), ("安全流转", "L4", "中")]),
        ("知识查询", [("本地知识库", "L3", "高"), ("制度检索", "L3", "高"), ("来源回溯", "L3", "中")]),
        ("沟通流转", [("通知汇总", "L4", "中"), ("消息草拟", "L2", "高"), ("自动发送", "L5", "低")]),
        ("数据分析", [("数据摘要", "L3", "中"), ("异常识别", "L4", "中"), ("趋势解读", "L4", "中")]),
        ("内容生成", [("文稿润色", "L2", "高"), ("汇报提纲", "L2", "高"), ("PPT/图表生成", "L4", "中")]),
    ],
}


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(FONT_BOLD if bold else FONT_REGULAR, size=size)


def centered(draw: ImageDraw.ImageDraw, xy: tuple[int, int], text: str, fnt, fill: str) -> None:
    box = draw.textbbox((0, 0), text, font=fnt)
    draw.text((xy[0] - (box[2] - box[0]) / 2, xy[1] - (box[3] - box[1]) / 2), text, font=fnt, fill=fill)


def rounded(draw: ImageDraw.ImageDraw, box, radius: int, fill: str, outline: str, width: int = 3) -> None:
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def distribute(count: int) -> list[int]:
    if count == 4:
        return [260, 575, 890, 1205]
    if count == 3:
        return [330, 730, 1130]
    if count == 2:
        return [470, 990]
    return [730]


def draw_connector(draw: ImageDraw.ImageDraw, points: Iterable[tuple[int, int]], color: str, width: int = 4) -> None:
    pts = list(points)
    draw.line(pts, fill=color, width=width, joint="curve")


def draw_leaf(draw: ImageDraw.ImageDraw, side: str, cy: int, item: tuple[str, str, str]) -> tuple[int, int]:
    name, difficulty, feasibility = item
    border, fill = FEASIBILITY[feasibility]
    if side == "left":
        box = (55, cy - 35, 500, cy + 35)
        anchor = (500, cy)
    else:
        box = (2100, cy - 35, 2545, cy + 35)
        anchor = (2100, cy)
    rounded(draw, box, 12, fill, border, 3)
    text_x = box[0] + 18
    draw.text((text_x, cy - 27), name, font=font(27, True), fill=INK)
    rating = f"{difficulty}  ·  可行性 {feasibility}"
    draw.text((text_x, cy + 5), rating, font=font(20), fill=border)
    return anchor


def draw_group(
    draw: ImageDraw.ImageDraw,
    center: tuple[int, int],
    side: str,
    title: str,
    items: list[tuple[str, str, str]],
    color_pair: tuple[str, str],
) -> None:
    color, fill = color_pair
    gx, gy = center
    group_box = (gx - 145, gy - 40, gx + 145, gy + 40)
    rounded(draw, group_box, 14, fill, color, 4)
    centered(draw, (gx, gy), title, font(29, True), color)

    offsets = [-84, 0, 84] if len(items) == 3 else [-45, 45]
    for item, offset in zip(items, offsets):
        leaf_y = gy + offset
        leaf_anchor = draw_leaf(draw, side, leaf_y, item)
        if side == "left":
            start = (gx - 145, gy)
            elbow_x = 545
            end = leaf_anchor
        else:
            start = (gx + 145, gy)
            elbow_x = 2055
            end = leaf_anchor
        draw_connector(draw, [start, (elbow_x, gy), (elbow_x, leaf_y), end], color, 3)


def draw_map(name: str, groups: list[tuple[str, list[tuple[str, str, str]]]]) -> Path:
    image = Image.new("RGB", CANVAS, BG)
    draw = ImageDraw.Draw(image)
    draw.text((70, 48), f"{name}功能思维导图", font=font(43, True), fill=INK)
    draw.text((72, 104), "功能节点仅标注实现难度与可行性", font=font(22), fill=MUTED)

    center = (1300, 760)
    center_box = (1135, 700, 1465, 820)
    rounded(draw, center_box, 18, CENTRAL, CENTRAL, 1)
    centered(draw, center, name, font(46, True), "#FFFFFF")

    left = groups[: (len(groups) + 1) // 2]
    right = groups[(len(groups) + 1) // 2 :]
    for side, side_groups in (("left", left), ("right", right)):
        ys = distribute(len(side_groups))
        for idx, ((group_name, items), gy) in enumerate(zip(side_groups, ys)):
            gx = 725 if side == "left" else 1875
            color_pair = GROUP_COLORS[(idx if side == "left" else idx + len(left)) % len(GROUP_COLORS)]
            color = color_pair[0]
            if side == "left":
                draw_connector(draw, [(1135, 760), (1010, 760), (1010, gy), (870, gy)], color, 5)
            else:
                draw_connector(draw, [(1465, 760), (1590, 760), (1590, gy), (1730, gy)], color, 5)
            draw_group(draw, (gx, gy), side, group_name, items, color_pair)

    out = ASSET_DIR / f"{name}.png"
    image.save(out, dpi=(220, 220), optimize=True)
    return out


def set_run_font(run, size: float, bold: bool = False, color: str = "17212B") -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def shade_paragraph(paragraph, fill: str, border: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "7")
    left.set(qn("w:color"), border)
    borders.append(left)


def set_image_alt(inline_shape, title: str, description: str) -> None:
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, 9, color="66727F")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_end])
    tail = paragraph.add_run(" 页")
    set_run_font(tail, 9, color="66727F")


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.15)
    section.bottom_margin = Cm(1.05)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)
    section.header_distance = Cm(0.45)
    section.footer_distance = Cm(0.45)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = doc.styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    title.font.size = Pt(28)
    title.font.bold = True
    title.font.color.rgb = RGBColor(38, 55, 70)
    title.paragraph_format.space_before = Pt(76)
    title.paragraph_format.space_after = Pt(8)

    subtitle = doc.styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    subtitle.font.size = Pt(19)
    subtitle.font.bold = True
    subtitle.font.color.rgb = RGBColor(46, 116, 181)
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(30)

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(46, 116, 181)
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(5)

    if "Map Caption" not in doc.styles:
        style = doc.styles.add_style("Map Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        style = doc.styles["Map Caption"]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(9)
    style.font.color.rgb = RGBColor(102, 114, 127)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(3)

    footer = section.footer
    p = footer.paragraphs[0]
    p.text = ""
    left = p.add_run("端云协同 AI 终端 · 四层功能思维导图")
    set_run_font(left, 9, color="66727F")
    p.add_run("\t")
    p.paragraph_format.tab_stops.add_tab_stop(Cm(25.5))
    add_page_number(p)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("端云协同 AI 终端")
    set_run_font(r, 28, True, "263746")

    p = doc.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("四层功能思维导图")
    set_run_font(r, 19, True, "2E74B5")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(22)
    r = p.add_run("架构层 · 应用层 · 生活层 · 工作层")
    set_run_font(r, 12, False, "66727F")

    h = doc.add_paragraph(style="Heading 1")
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.add_run("评级图例")

    labels = [
        ("结构", "层级 → 功能组 → 具体功能"),
        ("实现难度", "L1 配置集成 · L2 常规开发 · L3 多模块协同 · L4 复杂系统集成 · L5 高不确定性/生态依赖"),
        ("可行性 高", "技术成熟、实现边界清楚"),
        ("可行性 中", "可实现，但依赖适配或验证"),
        ("可行性 低", "存在明显平台、权限、可靠性或生态限制"),
    ]
    fills = ["E8EEF5", "E8EEF5", "EAF6F0", "FFF6E3", "FDECEF"]
    borders = ["2E74B5", "2E74B5", "2E7D5B", "B7791F", "C14953"]
    for (label, value), fill, border in zip(labels, fills, borders):
        line = doc.add_paragraph()
        line.paragraph_format.left_indent = Cm(2.2)
        line.paragraph_format.right_indent = Cm(2.2)
        line.paragraph_format.space_before = Pt(3)
        line.paragraph_format.space_after = Pt(3)
        shade_paragraph(line, fill, border)
        label_run = line.add_run(f"{label}  ")
        set_run_font(label_run, 10.5, True, "263746")
        value_run = line.add_run(value)
        set_run_font(value_run, 10.5, False, "263746")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run("内容依据现有产品设计方案与功能讨论整理")
    set_run_font(r, 9.5, False, "66727F")


def add_map_page(doc: Document, name: str, path: Path) -> None:
    heading = doc.add_paragraph(style="Heading 1")
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading.add_run(f"{name}功能思维导图")
    caption = doc.add_paragraph(style="Map Caption")
    caption.add_run("节点格式：功能名称｜实现难度｜可行性")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    inline = run.add_picture(str(path), width=Inches(10.55))
    set_image_alt(
        inline,
        f"{name}功能思维导图",
        f"{name}的两级功能树，包含功能组、具体功能、实现难度和可行性。",
    )


def build_docx(image_paths: dict[str, Path]) -> None:
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    doc.add_page_break()
    names = list(MAPS)
    for idx, name in enumerate(names):
        add_map_page(doc, name, image_paths[name])
        if idx < len(names) - 1:
            doc.add_page_break()
    doc.core_properties.title = "端云协同 AI 终端四层功能思维导图"
    doc.core_properties.subject = "架构层、应用层、生活层、工作层功能列表"
    doc.core_properties.keywords = "端云协同, AI终端, 思维导图, 功能列表"
    doc.save(OUTPUT)


def main() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    image_paths = {name: draw_map(name, groups) for name, groups in MAPS.items()}
    build_docx(image_paths)
    print(OUTPUT)
    for name, path in image_paths.items():
        print(name, path)


if __name__ == "__main__":
    main()
