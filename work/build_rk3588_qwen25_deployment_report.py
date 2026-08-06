from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\my new work\cloud flowing")
PACKAGE = ROOT / "端云协同AI终端产品功能设计-V0.7交付包"
OUTPUT = PACKAGE / "RK3588标准版-Qwen2.5-3B-Instruct端云功能部署评估报告.docx"

FONT_LATIN = "Calibri"
FONT_CJK = "Microsoft YaHei"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "243447"
MUTED = "667085"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "D0D5DD"
CALLOUT = "F4F6F9"
POSITIVE = "1F3A5F"
CAUTION = "7A5A00"
RISK = "9B1C1C"
WHITE = "FFFFFF"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_run_font(run, *, size=None, bold=None, italic=None, color=None,
                 latin=FONT_LATIN, cjk=FONT_CJK):
    run.font.name = latin
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    rfonts = run._element.rPr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        run._element.rPr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), cjk)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = rgb(color)


def set_style_font(style, *, size, color=INK, bold=False, latin=FONT_LATIN,
                   cjk=FONT_CJK):
    style.font.name = latin
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = rgb(color)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), cjk)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def set_no_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
    cant_split.set(qn("w:val"), "true")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    if sum(widths_dxa) != TABLE_WIDTH_DXA:
        raise ValueError(f"Table widths must sum to {TABLE_WIDTH_DXA}: {widths_dxa}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            set_cell_margins(cell)


def set_table_borders(table, color=MID_GRAY, size=4):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def paragraph_bottom_border(paragraph, color=BLUE, size=10, space=6):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)


def shade_paragraph(paragraph, fill=CALLOUT):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    left = borders.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        borders.append(left)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), BLUE)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)
    set_run_font(run, size=9, color=MUTED)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    relation_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), FONT_LATIN)
    rfonts.set(qn("w:hAnsi"), FONT_LATIN)
    rfonts.set(qn("w:eastAsia"), FONT_CJK)
    r_pr.append(rfonts)
    r_pr.append(color)
    r_pr.append(underline)
    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def configure_styles(doc):
    styles = doc.styles

    normal = styles["Normal"]
    set_style_font(normal, size=11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = styles["Title"]
    set_style_font(title, size=23, color="111827", bold=True)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)
    title.paragraph_format.keep_with_next = True

    subtitle = styles["Subtitle"]
    set_style_font(subtitle, size=14, color="374151")
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(16)
    subtitle.paragraph_format.keep_with_next = True

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        set_style_font(style, size=size, color=color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_name in ("List Bullet", "List Number"):
        style = styles[list_name]
        set_style_font(style, size=11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    if "Table Text" not in styles:
        table_text = styles.add_style("Table Text", WD_STYLE_TYPE.PARAGRAPH)
    else:
        table_text = styles["Table Text"]
    set_style_font(table_text, size=9.2)
    table_text.paragraph_format.space_before = Pt(0)
    table_text.paragraph_format.space_after = Pt(0)
    table_text.paragraph_format.line_spacing = 1.06

    if "Table Header" not in styles:
        table_header = styles.add_style("Table Header", WD_STYLE_TYPE.PARAGRAPH)
    else:
        table_header = styles["Table Header"]
    set_style_font(table_header, size=9.2, color="111827", bold=True)
    table_header.paragraph_format.space_before = Pt(0)
    table_header.paragraph_format.space_after = Pt(0)
    table_header.paragraph_format.line_spacing = 1.05

    caption = styles["Caption"]
    set_style_font(caption, size=9.5, color=MUTED, bold=True)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(4)
    caption.paragraph_format.keep_with_next = True

    if "Source Text" not in styles:
        source = styles.add_style("Source Text", WD_STYLE_TYPE.PARAGRAPH)
    else:
        source = styles["Source Text"]
    set_style_font(source, size=9.5, color=MUTED)
    source.paragraph_format.space_before = Pt(0)
    source.paragraph_format.space_after = Pt(4)
    source.paragraph_format.line_spacing = 1.05


def add_metadata_row(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.05
    label_run = p.add_run(f"{label}: ")
    set_run_font(label_run, size=10.5, bold=True, color="111827")
    value_run = p.add_run(value)
    set_run_font(value_run, size=10.5, color="344054")


def add_lead(doc, label, text, color=POSITIVE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.line_spacing = 1.10
    shade_paragraph(p)
    run = p.add_run(f"{label}  ")
    set_run_font(run, size=11, bold=True, color=color)
    run = p.add_run(text)
    set_run_font(run, size=11, color=INK)
    return p


def add_bullet(doc, text, *, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix and text.startswith(bold_prefix):
        run = p.add_run(bold_prefix)
        set_run_font(run, bold=True, color="111827")
        run = p.add_run(text[len(bold_prefix):])
        set_run_font(run, color=INK)
    else:
        run = p.add_run(text)
        set_run_font(run, color=INK)
    return p


def verdict_color(verdict):
    if verdict in {"本地闭环", "本地", "本地优先"}:
        return POSITIVE
    if verdict in {"端云协同", "本地有条件"}:
        return CAUTION
    return RISK


def add_matrix(doc, title, headers, rows, widths, *, verdict_col=None):
    cap = doc.add_paragraph(title, style="Caption")
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    set_no_split(table.rows[0])

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        shade_cell(cell, LIGHT_GRAY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.style = doc.styles["Table Header"]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_run_font(run, size=9.2, bold=True, color="111827")

    for row_data in rows:
        row = table.add_row()
        set_no_split(row)
        for index, value in enumerate(row_data):
            cell = row.cells[index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.style = doc.styles["Table Text"]
            if verdict_col is not None and index == verdict_col:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(value)
                set_run_font(run, size=9.2, bold=True, color=verdict_color(value))
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(value)
                set_run_font(run, size=9.2, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_step(doc, number, title, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    run = p.add_run(f"{number}. {title}  ")
    set_run_font(run, bold=True, color=DARK_BLUE)
    run = p.add_run(text)
    set_run_font(run, color=INK)


def add_source(doc, marker, title, url=None, note=None):
    p = doc.add_paragraph(style="Source Text")
    run = p.add_run(f"[{marker}] ")
    set_run_font(run, size=9.5, bold=True, color="111827")
    if url:
        add_hyperlink(p, title, url)
    else:
        run = p.add_run(title)
        set_run_font(run, size=9.5, color=INK)
    if note:
        run = p.add_run(f"。{note}")
        set_run_font(run, size=9.5, color=MUTED)


def build_document():
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    props = doc.core_properties
    props.title = "RK3588标准版-Qwen2.5-3B-Instruct端云功能部署评估报告"
    props.subject = "端云协同AI终端本地与云端功能边界评估"
    props.author = "Codex"
    props.keywords = "RK3588, Qwen2.5-3B-Instruct, RKLLM, W8A8, 端云协同"
    props.comments = "依据V0.7产品功能设计材料编制"

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    hr = hp.add_run("端云协同 AI 终端  |  技术评估")
    set_run_font(hr, size=8.5, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    fr = fp.add_run("V1.0  |  2026-07-28  |  ")
    set_run_font(fr, size=9, color=MUTED)
    add_page_field(fp)

    # Memo masthead
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(10)
    title = doc.add_paragraph(style="Title")
    title.add_run("RK3588标准版 + Qwen2.5-3B-Instruct").bold = True
    title2 = doc.add_paragraph(style="Title")
    title2.paragraph_format.space_before = Pt(0)
    title2.add_run("端云功能部署评估报告").bold = True
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("面向 V0.7 产品功能设计的本地、端云、联网服务与用户接管边界")

    add_metadata_row(doc, "评估对象", "RK3588 标准版硬件；Qwen2.5-3B-Instruct；RKLLM 量化部署")
    add_metadata_row(doc, "依据材料", "产品功能设计书、四层功能思维导图（修订版）、产品功能设计汇报稿")
    add_metadata_row(doc, "报告版本", "V1.0  |  评估日期 2026-07-28")
    add_metadata_row(doc, "决策状态", "有条件可行；先冻结 3-5 个高频工具并完成目标板验证")
    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(10)
    paragraph_bottom_border(rule)

    add_lead(
        doc,
        "核心结论",
        "RK3588 + Qwen2.5-3B-Instruct 适合做本地任务中枢，不适合被定义为全功能本地大模型终端。模型负责理解、参数提取、工具选择、短文本表达和状态汇报；确定性工具负责权限、执行、幂等、校验和审计。",
    )
    add_lead(
        doc,
        "首代范围",
        "文件与报表、本地知识、提醒待办、短文稿可本地优先；会议采用本地采集与端云转写/纪要；实时数据、复杂推理、长上下文和复杂内容生产默认使用云端或增强型平台。",
        CAUTION,
    )
    add_lead(
        doc,
        "硬件建议",
        "8 GB 仅作为串行负载样机；16 GB 更适合作为标准产品配置。两者都必须在目标板上验证模型转换、持续性能、内存峰值、散热降频以及 ASR/检索/UI 叠加负载。",
        CAUTION,
    )
    add_lead(
        doc,
        "安全边界",
        "下单、购票、支付、证券交易、删除和关键业务写回不因上云而变成可自动执行任务，必须使用授权接口或图形化接管，并保留用户最终确认。",
        RISK,
    )

    doc.add_heading("1. 评估口径", level=1)
    p = doc.add_paragraph()
    p.add_run("本报告回答的是产品级可交付性，而不是单次演示是否能跑通。").bold = True
    p.add_run("每项功能同时考虑模型能力、硬件资源、数据来源、工具可靠性、权限、安全、异常恢复和持续运行。")

    add_matrix(
        doc,
        "表 1  部署判定口径",
        ["判定", "含义", "产品承诺"],
        [
            ("本地闭环", "主要计算和执行在设备端完成，断网仍可使用。", "可作为标准版本地能力。"),
            ("本地优先", "常见短任务本地完成，超过上下文、质量或资源阈值后转云。", "必须公开云端触发条件。"),
            ("端云协同", "设备端保留身份、权限、缓存、状态和校验；云端承担部分计算或服务。", "不可宣传为完全离线。"),
            ("联网服务必需", "价值依赖实时数据、第三方平台或企业服务。", "断网时只能缓存、降级或明确失败。"),
            ("云端/增强型", "标准 RK3588 上不宜承诺稳定质量、长上下文或并发。", "转云或进入更高规格产品。"),
            ("用户接管", "涉及强认证、支付、交易或重大影响。", "模型仅准备与引导，不做无人值守提交。"),
        ],
        [1500, 4200, 3660],
        verdict_col=0,
    )

    doc.add_heading("2. 硬件与模型能力基线", level=1)
    p = doc.add_paragraph()
    p.add_run("已确认的型号为 Qwen2.5-3B-Instruct。").bold = True
    p.add_run("以下事实用于确定工程边界，但不把模型卡的理论上下文或其他 3B 模型的短上下文基准直接等同于本产品性能。")

    add_matrix(
        doc,
        "表 2  官方事实与工程解释",
        ["项目", "已核实事实", "工程解释"],
        [
            ("RK3588", "4x Cortex-A76 + 4x Cortex-A55；三核 6 TOPS NPU，支持 INT4/INT8/INT16/FP16 等。[S1]", "具备边缘推理基础，但系统能力仍受内存、带宽、散热和运行时影响。"),
            ("RKLLM", "支持 RK3588，并列出 Qwen2/Qwen2.5/Qwen3/Qwen3.5 模型家族；模型先在 PC 端转换/量化，再由板端 Runtime 推理。[S2]", "需要冻结 Toolkit、Runtime、驱动和固件版本，不能只冻结模型文件。"),
            ("Qwen2.5-3B-Instruct", "3.09B 参数、36 层，模型卡给出 32,768 token 上下文和 8,192 token 生成上限。[S4]", "这是模型能力上限，不是 RK3588 标准版的产品承诺；首版应从较短上下文开始实测。"),
            ("3B 参考量级", "Rockchip 短上下文基准中，TeleChat2 3B W8A8 在 RK3588 为 10.18 token/s、约 2,788 MB；不是 Qwen2.5-3B 的直接成绩。[S3]", "只能证明 3B 级量化模型可运行，不能证明长上下文、多模型并发或完整业务负载。"),
            ("Qwen2.5-VL-3B 参考", "官方多模态基准在 RK3588 W8A8 下解码约 8.66 token/s。[S3]", "多模态路径不同，仅作为负载量级参考；本报告不把视觉能力纳入标准版承诺。"),
            ("语音能力", "Qwen2.5-3B-Instruct 是因果语言模型，不是 ASR 或 TTS 模型。[S4]", "唤醒、语音识别和语音合成必须采用独立组件，并单独测量资源占用。"),
            ("许可", "模型页面标识为 qwen-research 许可。[S4]", "正式产品使用前必须由法务/采购核对许可范围；本文不作许可结论。"),
        ],
        [1700, 3900, 3760],
    )

    add_lead(
        doc,
        "关键判断",
        "模型卡支持 32K 上下文不等于 RK3588 产品可以稳定使用 32K。KV Cache、提示词预填、检索、语音和 UI 会共同占用内存与算力；必须依据目标负载设定本地上下文预算。",
        CAUTION,
    )

    doc.add_page_break()
    doc.add_heading("3. 用户功能端云部署矩阵", level=1)
    p = doc.add_paragraph()
    p.add_run("判定原则：").bold = True
    p.add_run("本地模型承担自然语言入口，不代替确定性服务、第三方数据源和业务授权。")

    feature_rows = [
        ("统一任务入口、状态、确认、取消", "本地闭环", "设备 UI、任务状态机、确认页和异常恢复。", "跨端查看时再接同步服务。"),
        ("授权文件查找与报表打开", "本地闭环", "授权索引、候选匹配、本地工具或 PC Agent、打开回执。", "跨企业系统时需要受控连接器。"),
        ("短文件摘要", "本地优先", "检索后片段或短文件交给 3B 模型摘要，并保留来源。", "文件过长、跨文档综合或质量不足时转云。"),
        ("长文档与跨文档综合", "云端/增强型", "本地负责切片、脱敏、上传范围确认和结果校验。", "复杂比较、长上下文和高质量综合由云端或增强型承担。"),
        ("本地知识与制度查询", "本地优先", "本地向量检索、权限过滤、短片段回答和来源引用。", "检索不足或复杂推理时，符合策略才可脱敏转云。"),
        ("提醒、日程与待办", "本地闭环", "时间解析、时区/重复规则、持久化、通知和状态更新。", "同步 Outlook、飞书等平台时需要授权 API。"),
        ("短文稿润色与消息草拟", "本地优先", "短文本改写、语气调整、草稿标识和事实字段保护。", "长报告、复杂结构或质量要求高时转云。"),
        ("复杂报告、方案和长篇内容", "云端/增强型", "本地负责材料选择、脱敏、任务状态和结果落盘。", "复杂规划、长链推理和长篇生成不作为标准版本地承诺。"),
        ("会议录音、加密缓存、停止/恢复", "本地闭环", "音频采集、录音状态、分片、加密缓存和完整性校验。", "原始音频是否上云由数据策略决定。"),
        ("近场短语音 ASR 与 TTS", "本地有条件", "使用独立 ASR/TTS/唤醒组件，可与 LLM 串行运行。", "不是 Qwen2.5-3B 功能；必须对中文、噪声和热稳态单测。"),
        ("长会议、多人远场转写与说话人区分", "端云协同", "本地完成 VAD、降噪、缓存、权限和进度。", "长时 ASR、远场识别和说话人分离建议云端或独立算力。"),
        ("会议纪要、行动项和来源回溯", "端云协同", "短转写可本地整理；保留原文、来源片段和草稿状态。", "长会议或复杂议题由云端处理，重要结论必须人工复核。"),
        ("手机 App/可携麦克风会议采集", "端云协同", "端侧录音、绑定、缓存、断点续传、去重和撤销信任。", "跨网络同步、设备管理和集中策略通常需要服务端。"),
        ("多端任务与通知同步", "端云协同", "设备端维护本地副本、版本和离线队列。", "跨网络同步需要云端或企业内网服务。"),
        ("本地应用、文件和白名单命令", "本地闭环", "确定性工具/PC Agent 执行，使用权限、幂等和结果校验。", "禁止模型无限制自由操作系统或不可预测页面。"),
        ("项目状态与业务记录查询", "端云协同", "本地解析请求、执行权限和统一展示。", "必须依赖 ERP/CRM/项目系统的稳定连接器和数据口径。"),
        ("外卖、航班、火车票、股票、天气资讯", "联网服务必需", "本地确认查询条件、最小化隐私、缓存并展示来源/时效。", "必须使用合法实时数据 API；不一定需要云大模型。"),
        ("下单、购票方案选择与预填", "端云协同", "本地组织候选方案、预览和确认页面。", "提交依赖平台 API 或图形化接管；价格变化后原确认失效。"),
        ("支付、实名购票、证券交易、删除", "用户接管", "设备显示对象、金额、影响、确认和审计。", "强认证在受信界面完成；禁止无人值守提交和保存关键凭据。"),
        ("BLE/Matter/局域网设备查询与控制", "端云协同", "兼容清单内可本地发现、状态读取、受控指令和回执。", "取决于无线模块、协议栈和设备适配；高风险控制必须确认。"),
        ("米家或其他厂商生态", "联网服务必需", "本地承担入口、缓存、权限、状态与失败反馈。", "优先使用官方授权云接口，不依赖私有协议绕行。"),
        ("经营概览、异常、风险和趋势", "云端/增强型", "终端承担交互、身份、显示、来源下钻和确认。", "依赖公司数据平台、指标口径、血缘、权限和质量治理。"),
        ("文档比较与结构化提取", "端云协同", "小文件、固定模板和规则抽取可本地完成。", "复杂版式、跨文档语义比较或大批量处理转云。"),
        ("PPT、图表和网页生成", "云端/增强型", "本地可做模板填充和文件交付。", "复杂设计、代码生成、素材处理和多轮修改不作为标准版本地能力。"),
        ("多指令拆分与并发执行", "云端/增强型", "标准版可拆分后排队，并优先保证前台确认、取消和状态。", "多个模型任务并行、长上下文和内容生成并发转云或增强型。"),
    ]
    add_matrix(
        doc,
        "表 3  用户功能部署判定",
        ["功能", "判定", "设备端职责", "云端/边界"],
        feature_rows,
        [2000, 1350, 3010, 3000],
        verdict_col=1,
    )

    doc.add_heading("4. 底层平台能力部署矩阵", level=1)
    platform_rows = [
        ("Qwen2.5-3B-Instruct 推理", "本地", "RKLLM W8A8 量化模型、板端 Runtime 和受控上下文。", "资源不足、校验失败或复杂任务转云。"),
        ("意图识别与参数提取", "本地", "模型输出结构化任务描述。", "所有输出必须经过 Schema 和业务规则校验。"),
        ("工具选择与结构化调用", "本地", "从注册白名单选择工具和参数。", "未注册工具、无校验器工具不得调用。"),
        ("歧义识别与补充确认", "本地", "展示候选对象、缺失参数和风险提示。", "模型置信度不能替代用户确认。"),
        ("工具注册、权限与风险策略", "本地", "规则、角色、数据域、设备信任和风险等级由确定性服务执行。", "模型不得授予、扩大或降低权限。"),
        ("工具执行、幂等与结果校验", "本地", "生成幂等键、记录回执、验证对象和状态。", "外部写操作状态不明时先查实，不盲目重试。"),
        ("端云路由判定", "本地", "综合工具可用性、数据级别、资源、网络和模型校验结果。", "路由原因和实际出端数据范围必须可见。"),
        ("数据分级与脱敏", "本地优先", "字段级规则、敏感词/模式、策略和禁止外发清单。", "语义识别可辅助，但 D3 凭据不得进入模型上下文。"),
        ("任务状态机与异常恢复", "本地", "保存阶段、取消、重试预算、部分成功和恢复选项。", "云端失败不得无限循环切换。"),
        ("离线任务队列", "本地", "断网保留上下文和待同步状态。", "恢复网络后按版本、幂等和权限继续。"),
        ("全链路审计", "本地优先", "本地记录任务、路由、确认、工具版本、参数摘要和回执。", "组织级审计可加密汇聚到云端或内网平台。"),
        ("账号与设备绑定", "端云协同", "设备保存可信状态和本地密钥。", "跨设备身份、撤销和组织管理需要服务端。"),
        ("传输加密与凭据管理", "端云协同", "本地安全存储、最小权限和短期凭据。", "证书、密钥轮换和统一身份由企业/云服务配合。"),
        ("模型、提示和工具版本管理", "端云协同", "板端版本锁定、兼容检查和回滚。", "灰度、制品分发和策略运营需要管理服务。"),
        ("监控与远程升级", "端云协同", "设备健康、日志缓冲、升级校验和本地回滚。", "集中监控、告警和升级编排需要云端/内网平台。"),
        ("ASR/TTS/唤醒运行时", "本地有条件", "作为独立模型和服务部署，设置独立资源配额。", "会议级语音需云端或单独的本地性能方案。"),
        ("资源调度", "本地", "前台优先、队列、超时、取消、内存水位和资源回收。", "标准版默认单模型推理，不承诺多个模型任务并发。"),
        ("第三方与企业连接器", "端云协同", "本地保存最小配置并做请求校验。", "授权、API 可用性、速率限制和数据质量由上游决定。"),
    ]
    add_matrix(
        doc,
        "表 4  平台能力部署判定",
        ["平台能力", "判定", "本地实现", "约束/服务端职责"],
        platform_rows,
        [2100, 1300, 2980, 2980],
        verdict_col=1,
    )

    doc.add_heading("5. 推荐的端云任务链", level=1)
    add_step(doc, 1, "本地接收", "语音或文字输入进入设备端，立即创建任务编号并显示已接收状态。")
    add_step(doc, 2, "本地理解", "Qwen2.5-3B-Instruct 输出意图、对象、动作、时间、数据范围、输出形式和缺失参数。")
    add_step(doc, 3, "确定性校验", "Schema、权限、对象候选、风险等级、数据分级和本地资源水位共同决定是否可执行。")
    add_step(doc, 4, "路由", "低风险且有本地工具时本地执行；实时数据、复杂推理、长上下文或校验失败时按策略转云；禁止外发的数据不得转云。")
    add_step(doc, 5, "确认或接管", "涉及外部影响时展示对象、内容、金额、影响和可撤销性；R3 动作进入强认证或图形化接管。")
    add_step(doc, 6, "执行与校验", "工具产生真实回执后再宣告成功；状态不明时查询实际状态，避免重复提交。")
    add_step(doc, 7, "交付与审计", "结果显示来源、更新时间、执行端、完成状态和后续操作，并记录实际出端数据摘要。")

    add_lead(
        doc,
        "不可妥协",
        "模型只负责理解和表达，不直接成为权限引擎、支付执行器、证券交易代理或业务成功判定器。",
        RISK,
    )

    doc.add_heading("6. 标准版硬件配置建议", level=1)
    add_matrix(
        doc,
        "表 5  内存档位与能力边界",
        ["配置", "定位", "可承诺能力", "限制"],
        [
            ("RK3588 + 8 GB", "串行 PoC/样机", "W8A8 3B 模型、UI、轻量检索和单个工具任务；ASR 与 LLM 尽量错峰。", "不承诺长上下文、多模型常驻、会议级语音或多任务并发。"),
            ("RK3588 + 16 GB", "推荐标准版", "3B 模型、Embedding/Reranker、任务服务、UI、本地数据库及一个后台工具具有更可信余量。", "复杂报告、长会议、长上下文和重并发仍应转云。"),
            ("更高内存/算力平台", "增强型", "复杂本机内容、更多上下文和多任务调度。", "是否立项以标准版真实瓶颈和新增价值为前提。"),
        ],
        [1750, 1700, 3200, 2710],
    )

    add_bullet(doc, "内存：首版产品优先 16 GB；任何档位都保持至少 20% 运行余量，并使用内存水位触发排队或转云。", bold_prefix="内存：")
    add_bullet(doc, "上下文：从约 4K 的产品预算开始验证，采用检索切片和摘要压缩；扩展前必须测量 KV Cache 与预填时延。", bold_prefix="上下文：")
    add_bullet(doc, "存储：建议预留模型、索引、日志、音频缓存和回滚镜像空间，并为会议数据设置独立配额与自动清理。", bold_prefix="存储：")
    add_bullet(doc, "散热：使用产品外壳、目标供电和环境温度完成至少 60 分钟持续负载测试，记录降频后的稳定吞吐。", bold_prefix="散热：")
    add_bullet(doc, "版本：冻结 RKLLM Toolkit、Runtime、NPU 驱动、系统镜像、量化参数、Tokenizer 和 Chat Template 的兼容组合。", bold_prefix="版本：")

    doc.add_page_break()
    doc.add_heading("7. 首代 MVP 建议", level=1)
    p = doc.add_paragraph()
    p.add_run("建议首批冻结 4 个完整本地工具 + 1 个端云会议工具。").bold = True
    p.add_run("这样既能验证本地价值，也能验证路由、脱敏、云端失败和数据不允许出端时的降级能力。")

    add_matrix(
        doc,
        "表 6  推荐首批工具与交付顺序",
        ["顺序", "工具", "默认执行端", "首版验收重点"],
        [
            ("1", "授权文件查找与报表打开", "本地", "不越权；同名文件确认；打开成功有真实回执。"),
            ("2", "本地知识与制度查询", "本地优先", "关键结论有来源；无可靠依据时不编造；版本和权限正确。"),
            ("3", "提醒、日程与待办", "本地", "时间/重复规则正确；修改对象明确；取消后不再通知。"),
            ("4", "短文稿润色与消息草拟", "本地优先", "保留事实字段；标识未发送；发送前独立确认。"),
            ("5", "会议记录与纪要", "端云协同", "本地采集不丢失；转写质量可测；纪要可回溯；数据上云受策略控制。"),
        ],
        [900, 2600, 1600, 4260],
    )

    doc.add_heading("8. 样机验收门槛", level=1)
    p = doc.add_paragraph()
    p.add_run("以下是建议的起始门槛，不是已经达成的结果。").bold = True
    p.add_run("项目应在目标板首轮基线后冻结正式数值，并按场景分别统计 P50/P95 和长尾失败。")

    acceptance_rows = [
        ("模型部署", "Qwen2.5-3B-Instruct 完成 RKLLM W8A8 转换；固定版本可重复构建；Tokenizer/模板一致。"),
        ("持续性能", "512 token 级输入下建议 P95 首字时间不高于 2.5 秒，热稳态持续解码不低于 8 token/s；以目标板实测修订。"),
        ("资源", "60 分钟组合负载无 OOM、崩溃或不可恢复降频；峰值总内存不超过物理内存的 80%。"),
        ("工具编排", "意图 Top-1 准确率建议不低于 95%，必填参数准确率不低于 98%；未注册工具执行为 0。"),
        ("权限与风险", "越权数据泄露为 0；高风险动作确认覆盖率 100%；D3 凭据进入模型上下文或外发为 0。"),
        ("本地知识", "组织事实类回答的来源覆盖率 100%；无来源时必须拒绝确定性回答；权限隔离用例全部通过。"),
        ("异常恢复", "应用重启、断网、工具超时、云端失败和重复请求均有明确状态；写操作不得重复提交。"),
        ("前台响应", "一个后台工具任务运行时，状态展示、确认和取消保持可用，不被模型推理阻塞。"),
        ("语音", "若承诺本地 ASR，应使用定义好的中文近场/噪声测试集冻结字错率；远场和多人会议单独验收。"),
        ("端云透明", "每次转云均可说明原因、发送范围和执行端；禁止上云的数据任务不得上传。"),
    ]
    add_matrix(
        doc,
        "表 7  建议验收门槛",
        ["维度", "建议起始门槛"],
        acceptance_rows,
        [1900, 7460],
    )

    doc.add_page_break()
    doc.add_heading("9. 立项前必须冻结的决策", level=1)
    decisions = [
        "标准版具体内存、存储、散热、供电和系统镜像，而不仅是 RK3588 芯片名称。",
        "Qwen2.5-3B-Instruct 的量化方案、RKLLM 版本组合、上下文预算和产品许可结论。",
        "本地 ASR/TTS/唤醒组件，以及近场、会议室和移动采集分别采用的技术路线。",
        "首批 3-5 个工具、每个工具的权限、输入 Schema、成功判定、幂等、超时和撤销规则。",
        "D0-D3 数据分级、允许上云范围、脱敏方式、云端供应方、留存期限和审计要求。",
        "实时信息、票务、外卖、股票和企业系统的合法接口、授权、成本、时效及故障降级。",
        "标准版并发目标：前台交互之外允许几个后台工具、是否允许 ASR 与 LLM 同时常驻。",
    ]
    for item in decisions:
        add_bullet(doc, item)

    add_lead(
        doc,
        "最终建议",
        "以 16 GB RK3588 标准版作为产品样机基线，采用 Qwen2.5-3B-Instruct W8A8 做本地理解与编排，先交付 4 个本地工具和 1 个端云会议工具。通过目标板性能、安全和闭环验收后，再决定是否扩大本地上下文、语音负载或启动增强型。",
        POSITIVE,
    )

    doc.add_heading("10. 参考资料", level=1)
    add_source(doc, "D1", "端云协同AI终端产品功能设计书.md", note="V0.7 交付包内源文件")
    add_source(doc, "D2", "端云协同AI终端四层功能思维导图-修订版.md", note="V0.7 交付包内源文件")
    add_source(doc, "D3", "端云协同AI终端产品功能设计汇报稿.docx", note="V0.7 交付包内源文件")
    add_source(doc, "S1", "Rockchip RK3588 产品规格", "https://www.rock-chips.com/a/en/products/RK35_Series/2022/0926/1660.html", "访问日期 2026-07-28")
    add_source(doc, "S2", "Rockchip RKLLM 官方仓库", "https://github.com/airockchip/rknn-llm", "访问日期 2026-07-28")
    add_source(doc, "S3", "Rockchip RKLLM 模型性能基准", "https://github.com/airockchip/rknn-llm/blob/main/benchmark.md", "访问日期 2026-07-28")
    add_source(doc, "S4", "Qwen2.5-3B-Instruct 官方模型卡", "https://huggingface.co/Qwen/Qwen2.5-3B-Instruct", "访问日期 2026-07-28")
    add_source(doc, "S5", "Qwen2.5 LLM 官方技术说明", "https://qwenlm.github.io/blog/qwen2.5-llm/", "访问日期 2026-07-28")

    p = doc.add_paragraph(style="Source Text")
    r = p.add_run("说明：")
    set_run_font(r, size=9.5, bold=True, color="111827")
    r = p.add_run("性能数字仅作为官方参考量级；其中 3B W8A8 结果并非 Qwen2.5-3B-Instruct 文本模型的直接板测成绩。最终产品边界必须以目标板、最终量化模型和完整业务负载实测为准。")
    set_run_font(r, size=9.5, color=MUTED)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
